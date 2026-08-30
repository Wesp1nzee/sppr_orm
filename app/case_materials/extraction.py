"""Извлечение текста и сегментация загруженных файлов материалов дела.

Слой не зависит от БД и FastAPI: принимает байты файла, возвращает
``ExtractedText`` (полный текст + страницы) и список ``DetectedDocument``.
Точка расширения для будущего OCR — протокол ``TextExtractor``.
"""

import re
from dataclasses import dataclass, field
from typing import Any, Protocol

from app.case_materials.constants import DOCUMENT_KIND_TITLES, CaseDocumentKind
from app.case_materials.fields import FIELD_EXTRACTORS, ExtractedField
from app.case_materials.fields.base import FieldExtractionResult, FieldExtractor

PDF_MIME_TYPE = "application/pdf"
DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


@dataclass
class ExtractedText:
    """Результат извлечения текста: полный текст и текст по страницам.

    Для DOCX «страниц» нет — ``pages`` содержит один элемент (весь текст).
    """

    text: str
    pages: list[str] = field(default_factory=list)

    def page_for_offset(self, offset: int) -> int | None:
        """Номер страницы (1-based) для смещения символа; None для DOCX."""
        if len(self.pages) <= 1:
            return None
        cursor = 0
        for index, page in enumerate(self.pages, start=1):
            cursor += len(page)
            if offset < cursor:
                return index
        return len(self.pages)


class TextExtractor(Protocol):
    """Извлекает текст из бинарного содержимого файла."""

    def extract(self, data: bytes, filename: str) -> ExtractedText: ...


class PdfTextExtractor:
    """Извлечение текста из цифрового PDF через ``pypdf`` (текстовый слой)."""

    def extract(self, data: bytes, filename: str) -> ExtractedText:
        del filename
        import io

        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
        return ExtractedText(text="\n".join(pages), pages=pages)


class DocxTextExtractor:
    """Извлечение текста из DOCX через ``python-docx`` (абзацы + таблицы)."""

    def extract(self, data: bytes, filename: str) -> ExtractedText:
        del filename
        import io

        from docx import Document

        document = Document(io.BytesIO(data))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text:
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        text = "\n".join(parts)
        return ExtractedText(text=text, pages=[text])


@dataclass
class DetectedDocument:
    """Один распознанный документ внутри загруженного файла."""

    kind: CaseDocumentKind
    title: str
    start: int
    end: int
    page: int | None
    fields: dict[str, ExtractedField] = field(default_factory=dict)
    suggested_answers: dict[str, dict[str, Any]] = field(default_factory=dict)


def extractor_for(mime_type: str) -> TextExtractor:
    """Возвращает экстрактор по MIME-типу; для неизвестного — заглушку None-типа."""
    if mime_type == PDF_MIME_TYPE:
        return PdfTextExtractor()
    if mime_type == DOCX_MIME_TYPE:
        return DocxTextExtractor()
    raise ValueError(f"Unsupported mime type: {mime_type}")


def _header_patterns() -> list[tuple[CaseDocumentKind, re.Pattern[str]]]:
    """Паттерны заголовков документов (по референсному PDF, порядок важен)."""
    return [
        (
            CaseDocumentKind.resolution_to_conduct_orm,
            re.compile(
                r"ПОСТАНОВЛЕНИЕ\s*\n\s*о\s+проведении\s+оперативно-розыскного\s+мероприятия",
                re.IGNORECASE,
            ),
        ),
        (
            CaseDocumentKind.declassification_resolution,
            re.compile(r"ПОСТАНОВЛЕНИЕ\s*\n\s*о\s+рассекречивании", re.IGNORECASE),
        ),
        (
            CaseDocumentKind.representation_resolution,
            re.compile(
                r"ПОСТАНОВЛЕНИЕ\s*\n\s*о\s+представлении\s+результатов", re.IGNORECASE
            ),
        ),
        (
            CaseDocumentKind.inspection_order,
            re.compile(r"РАСПОРЯЖЕНИЕ\s+№", re.IGNORECASE),
        ),
        (
            CaseDocumentKind.consent_statement,
            re.compile(r"ЗАЯВЛЕНИЕ\s*\n", re.IGNORECASE),
        ),
        (
            CaseDocumentKind.search_protocol_before,
            re.compile(
                r"ПРОТОКОЛ\s*\n\s*досмотра\s+лица,\s*участвующего", re.IGNORECASE
            ),
        ),
        (
            CaseDocumentKind.search_protocol_after,
            re.compile(
                r"ПРОТОКОЛ\s*\n\s*досмотра\s+лица,\s*участвовавшего", re.IGNORECASE
            ),
        ),
        (
            CaseDocumentKind.money_inspection_protocol,
            re.compile(
                r"ПРОТОКОЛ\s*\n\s*осмотра,\s*пометки\s+и\s+вручения", re.IGNORECASE
            ),
        ),
        (
            CaseDocumentKind.sts_delivery_protocol,
            re.compile(r"ПРОТОКОЛ\s*\n\s*вручения\s+специального", re.IGNORECASE),
        ),
        (
            CaseDocumentKind.sts_withdrawal_protocol,
            re.compile(r"ПРОТОКОЛ\s*\n\s*изъятия\s+специального", re.IGNORECASE),
        ),
        (
            CaseDocumentKind.vehicle_inspection_protocol_before,
            re.compile(
                r"ПРОТОКОЛ\s*\n\s*обследования\s+автомототранспортного\s+средства\s+до",
                re.IGNORECASE,
            ),
        ),
        (
            CaseDocumentKind.vehicle_inspection_protocol_after,
            re.compile(
                r"ПРОТОКОЛ\s*\n\s*обследования\s+автомототранспортного\s+средства\s+после",
                re.IGNORECASE,
            ),
        ),
        (
            CaseDocumentKind.raport,
            re.compile(r"РАПОРТ\s*\n\s*о\s+проведении", re.IGNORECASE),
        ),
    ]


_HEADER_PATTERNS = _header_patterns()


class DocumentSegmenter:
    """Разбивает извлечённый текст на сегменты по заголовкам документов."""

    def segment(self, extracted: ExtractedText) -> list[DetectedDocument]:
        matches: list[tuple[int, int, CaseDocumentKind]] = []
        for kind, pattern in _HEADER_PATTERNS:
            for match in pattern.finditer(extracted.text):
                matches.append((match.start(), match.end(), kind))
        matches.sort(key=lambda item: item[0])

        documents: list[DetectedDocument] = []
        for index, (start, _, kind) in enumerate(matches):
            end = (
                matches[index + 1][0]
                if index + 1 < len(matches)
                else len(extracted.text)
            )
            result = self._extract(kind, extracted.text[start:end])
            documents.append(
                DetectedDocument(
                    kind=kind,
                    title=DOCUMENT_KIND_TITLES[kind],
                    start=start,
                    end=end,
                    page=extracted.page_for_offset(start),
                    fields=result.fields,
                    suggested_answers=result.suggested_answers,
                )
            )
        return documents

    @staticmethod
    def _extract(kind: CaseDocumentKind, text: str) -> FieldExtractionResult:
        extractor: FieldExtractor | None = FIELD_EXTRACTORS.get(kind)
        if extractor is None:
            return FieldExtractionResult()
        return extractor.extract(text)


__all__ = [
    "DetectedDocument",
    "DocumentSegmenter",
    "DocxTextExtractor",
    "ExtractedText",
    "PdfTextExtractor",
    "TextExtractor",
    "extractor_for",
]
