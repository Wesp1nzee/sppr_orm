"""Экспорт сгенерированных документов в DOCX и PDF (ТЗ, раздел 3.4).

DOCX строится через ``python-docx``; PDF — через ``reportlab`` (чистый Python,
без системной зависимости LibreOffice). Оба рендера идут из одного источника —
структурированного ``GeneratedDocument.content``.
"""

import html
import io
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4  # type: ignore[import-untyped]
from reportlab.lib.styles import (  # type: ignore[import-untyped]
    ParagraphStyle,
    getSampleStyleSheet,
)
from reportlab.lib.units import mm  # type: ignore[import-untyped]
from reportlab.pdfbase import pdfmetrics  # type: ignore[import-untyped]
from reportlab.pdfbase.ttfonts import TTFont  # type: ignore[import-untyped]
from reportlab.platypus import (  # type: ignore[import-untyped]
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

#: Заголовки разделов для отображения в файле.
SECTION_TITLES: dict[str, str] = {
    "адресат": "Адресат",
    "данные_заявителя": "Данные заявителя",
    "обстоятельства_дела": "Обстоятельства дела",
    "перечень_нарушений": "Перечень нарушений",
    "ссылки_на_кс_рф": "Ссылки на определения Конституционного Суда РФ",
    "правовое_обоснование": "Правовое обоснование",
    "просительная_часть": "Просительная часть",
    "дата": "Дата",
    "подпись": "Подпись",
    "заголовок": "Заголовок",
    "пункты": "Пункты",
    "этапы": "Этапы",
}

#: Подписи полей внутри элементов списков (нарушения, пункты чек-листа и т.п.).
FIELD_LABELS: dict[str, str] = {
    "criterion_number": "Критерий",
    "title": "Наименование",
    "comment": "Комментарий",
    "references": "Нормы",
    "number": "№",
    "summary": "Содержание",
    "номер": "Этап",
    "действие": "Действие",
    "комментарий": "Комментарий",
    "status": "Статус",
    "recommendations": "Рекомендации",
}


def export_docx(content: dict[str, Any], title: str) -> bytes:
    """Собирает DOCX-файл из структурированного содержимого документа."""
    doc = Document()
    doc.add_heading(title, 0)
    for key, value in content.items():
        doc.add_heading(SECTION_TITLES.get(key, key), level=1)
        _render_docx_value(doc, value)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_pdf(content: dict[str, Any], title: str) -> bytes:
    """Собирает PDF через reportlab (DejaVu для кириллицы, fallback — Helvetica)."""
    font_name = _register_font()
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4, topMargin=20 * mm, bottomMargin=20 * mm
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(name="DocTitle", fontName=font_name, fontSize=16, leading=20)
    )
    styles.add(
        ParagraphStyle(name="DocSection", fontName=font_name, fontSize=13, leading=17)
    )
    styles.add(
        ParagraphStyle(name="DocBody", fontName=font_name, fontSize=11, leading=15)
    )
    styles.add(
        ParagraphStyle(name="DocBullet", fontName=font_name, fontSize=11, leading=15)
    )

    story: list[Any] = [Paragraph(_esc(title), styles["DocTitle"]), Spacer(1, 6 * mm)]
    for key, value in content.items():
        story.append(
            Paragraph(_esc(SECTION_TITLES.get(key, key)), styles["DocSection"])
        )
        story.extend(_pdf_blocks(value, styles))
    doc.build(story)
    return buffer.getvalue()


def _render_docx_value(doc: Any, value: Any) -> None:
    if isinstance(value, str):
        for line in value.split("\n"):
            doc.add_paragraph(line)
    elif isinstance(value, list):
        for item in value:
            doc.add_paragraph(_format_item(item), style="List Bullet")
    elif isinstance(value, dict):
        for sub_key, sub_value in value.items():
            key = str(sub_key)
            doc.add_paragraph(SECTION_TITLES.get(key, key), style="List Bullet")
            _render_docx_value(doc, sub_value)
    else:
        doc.add_paragraph(str(value))


def _pdf_blocks(value: Any, styles: dict[str, ParagraphStyle]) -> list[Any]:
    blocks: list[Any] = []
    if isinstance(value, str):
        for line in value.split("\n"):
            blocks.append(Paragraph(_esc(line), styles["DocBody"]))
    elif isinstance(value, list):
        for item in value:
            blocks.append(Paragraph(_esc(_format_item(item)), styles["DocBullet"]))
    elif isinstance(value, dict):
        for sub_key, sub_value in value.items():
            key = str(sub_key)
            blocks.append(
                Paragraph(_esc(SECTION_TITLES.get(key, key)), styles["DocBullet"])
            )
            blocks.extend(_pdf_blocks(sub_value, styles))
    else:
        blocks.append(Paragraph(_esc(str(value)), styles["DocBody"]))
    return blocks


def _format_item(item: Any) -> str:
    """Форматирует элемент списка (строку или словарь) в читаемый текст."""
    if isinstance(item, str):
        return item
    if isinstance(item, dict):
        parts: list[str] = []
        for key, value in item.items():
            if value in (None, "", [], {}):
                continue
            if isinstance(value, list):
                value = ", ".join(str(part) for part in value)
            parts.append(f"{FIELD_LABELS.get(key, key)}: {value}")
        return "; ".join(parts)
    return str(item)


def _esc(text: str) -> str:
    return html.escape(text)


_FONT_PATHS = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
]


def _register_font() -> str:
    """Регистрирует DejaVu (кириллица); при отсутствии — стандартный Helvetica."""
    for path in _FONT_PATHS:
        if Path(path).exists():
            pdfmetrics.registerFont(TTFont("DejaVu", path))
            return "DejaVu"
    return "Helvetica"


__all__ = ["export_docx", "export_pdf"]
